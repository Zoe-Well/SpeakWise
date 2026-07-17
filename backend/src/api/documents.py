"""文档导入与素材附加 API"""

import json
import io
import asyncio
import logging
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Body
from fastapi.responses import StreamingResponse
from sqlmodel import Session

from backend.src.db.connection import get_session
from backend.src.services import profile_service
from backend.src.models.document import SourceDocument, ProfileUpdateProposal

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["documents"])


@router.post("/documents")
async def upload_document(
    file: UploadFile = File(...),
    scope: str = Form("profile"),
    usage: str = Form("parse"),
    clear_existing: bool = Form(False),
    session: Session = Depends(get_session),
):
    profile = profile_service.get_active_profile(session)
    filename = file.filename or "unknown"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    raw = await file.read()
    if len(raw) > 10 * 1024 * 1024:  # 10 MB limit
        raise HTTPException(413, "文件过大，最大支持 10MB")
    extracted = ""
    if ext in ("txt", "md"):
        extracted = raw.decode("utf-8", errors="ignore")
    elif ext == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            extracted = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            pass
    elif ext in ("docx", "doc"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            extracted = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            pass

    parse_status = "success" if extracted.strip() else "failed"

    doc = SourceDocument(profile_id=profile.id, scope=scope, usage=usage,
                         filename=filename, file_type=ext, extracted_text=extracted,
                         parse_status=parse_status)
    session.add(doc)
    session.commit()
    session.refresh(doc)

    result = {"document_id": doc.id, "parse_status": parse_status, "filename": filename}

    if usage == "parse" and extracted.strip():
        if scope == "profile":
            try:
                proposals = await _llm_parse_resume(extracted)
            except Exception:
                logger.exception("LLM 解析调用失败: doc_id=%d, filename=%s", doc.id, filename)
                proposals = []
                result["parse_error"] = "LLM 服务调用失败，请稍后重试"

            if proposals:
                prop = ProfileUpdateProposal(
                    profile_id=profile.id, document_id=doc.id,
                    changes=json.dumps(proposals, ensure_ascii=False), status="pending",
                    clear_existing=clear_existing,
                )
                session.add(prop)
                session.commit()
                session.refresh(prop)
                result["proposal"] = {"proposal_id": prop.id, "changes": proposals, "clear_existing": clear_existing}
            elif "parse_error" not in result:
                logger.info("LLM 解析未提取到字段: doc_id=%d, text_len=%d, filename=%s",
                            doc.id, len(extracted), filename)
                result["parse_detail"] = "no_fields"
        elif scope == "jd":
            # For JD docs: return extracted text so frontend can populate textarea
            result["extracted_text"] = extracted[:10000]

    return result


@router.get("/documents")
def list_documents(scope: str = None, session: Session = Depends(get_session)):
    """列出文档。scope 可选：不传返回所有，传 'profile' 或 'jd' 过滤。"""
    profile = profile_service.get_active_profile(session)
    from sqlmodel import select
    stmt = select(SourceDocument).where(SourceDocument.profile_id == profile.id)
    if scope:
        stmt = stmt.where(SourceDocument.scope == scope)
    docs = session.exec(stmt.order_by(SourceDocument.created_at.desc())).all()
    return [{"id": d.id, "filename": d.filename, "scope": d.scope, "usage": d.usage,
             "file_type": d.file_type, "parse_status": d.parse_status,
             "created_at": d.created_at.isoformat()} for d in docs]


@router.delete("/documents/{doc_id}")
def delete_document(doc_id: int, session: Session = Depends(get_session)):
    profile = profile_service.get_active_profile(session)
    doc = session.get(SourceDocument, doc_id)
    if not doc or doc.profile_id != profile.id:
        raise HTTPException(404)
    session.delete(doc)
    session.commit()
    return {"ok": True}


@router.post("/documents/{doc_id}/reparse")
async def reparse_document(doc_id: int, session: Session = Depends(get_session)):
    """对已有文档重新调用 LLM 解析（用于解析失败或结果不理想的文档）。"""
    profile = profile_service.get_active_profile(session)
    doc = session.get(SourceDocument, doc_id)
    if not doc or doc.profile_id != profile.id:
        raise HTTPException(404)
    if not doc.extracted_text or not doc.extracted_text.strip():
        raise HTTPException(400, "文档无已提取文本，无法重新解析")

    logger.info("重新解析文档 doc_id=%d, filename=%s", doc.id, doc.filename)
    proposals = await _llm_parse_resume(doc.extracted_text)

    # 将旧 pending proposals 标记为 rejected
    from sqlmodel import select as _sel
    old_props = session.exec(_sel(ProfileUpdateProposal).where(
        ProfileUpdateProposal.document_id == doc_id,
        ProfileUpdateProposal.status == "pending"
    )).all()
    for op in old_props:
        op.status = "rejected"
        session.add(op)

    if proposals:
        prop = ProfileUpdateProposal(
            profile_id=profile.id, document_id=doc.id,
            changes=json.dumps(proposals, ensure_ascii=False), status="pending",
        )
        session.add(prop)
        session.commit()
        session.refresh(prop)
        return {"proposal": {"proposal_id": prop.id, "changes": proposals}}
    else:
        session.commit()
        logger.info("重新解析未提取到字段: doc_id=%d", doc.id)
        return {"proposal": None, "message": "未提取到新字段"}


@router.post("/documents/batch-delete")
def batch_delete_documents(data: dict = Body(...), session: Session = Depends(get_session)):
    """批量删除文档。body: { ids: [1, 2, 3] }"""
    profile = profile_service.get_active_profile(session)
    ids = data.get("ids", [])
    deleted = 0
    for doc_id in ids:
        doc = session.get(SourceDocument, doc_id)
        if doc and doc.profile_id == profile.id:
            session.delete(doc)
            deleted += 1
    session.commit()
    return {"deleted": deleted}


# ── Material Selection (for Knowledge Selector) ──────────

@router.get("/materials")
def list_materials(scope: str = "profile", usage: str = "attach", session: Session = Depends(get_session)):
    """列出可选择的素材文档（按 scope + usage 过滤）。"""
    profile = profile_service.get_active_profile(session)
    from sqlmodel import select
    docs = session.exec(
        select(SourceDocument)
        .where(SourceDocument.profile_id == profile.id)
        .where(SourceDocument.scope == scope)
        .where(SourceDocument.usage == usage)
        .order_by(SourceDocument.created_at.desc())
    ).all()
    return [{
        "id": d.id, "filename": d.filename, "scope": d.scope, "usage": d.usage,
        "file_type": d.file_type, "is_active": d.is_active,
        "created_at": d.created_at.isoformat() if d.created_at else None,
    } for d in docs]


@router.put("/materials/{doc_id}/toggle")
def toggle_material(doc_id: int, session: Session = Depends(get_session)):
    """切换素材的 is_active 状态（多选）。"""
    profile = profile_service.get_active_profile(session)
    doc = session.get(SourceDocument, doc_id)
    if not doc or doc.profile_id != profile.id:
        raise HTTPException(404, "文档不存在")
    doc.is_active = not doc.is_active
    session.add(doc)
    session.commit()
    return {"ok": True, "id": doc.id, "is_active": doc.is_active}


@router.get("/profile/proposals")
def list_pending_proposals(session: Session = Depends(get_session)):
    """列出当前 profile 的待处理简历解析建议。"""
    profile = profile_service.get_active_profile(session)
    from sqlmodel import select as _sel
    rows = session.exec(_sel(ProfileUpdateProposal).where(
        ProfileUpdateProposal.profile_id == profile.id,
        ProfileUpdateProposal.status == "pending"
    ).order_by(ProfileUpdateProposal.created_at.desc())).all()
    return [{"proposal_id": r.id, "document_id": r.document_id,
             "changes": json.loads(r.changes) if r.changes else [],
             "status": r.status, "clear_existing": r.clear_existing} for r in rows]


@router.post("/profile/merge")
def merge_proposal(data: dict, session: Session = Depends(get_session)):
    """用户确认写入知识库更新建议。仅写入 accepted_change_ids 中的变更。"""
    proposal_id = data.get("proposal_id")
    accepted_ids = set(data.get("accepted_change_ids", []))
    prop = session.get(ProfileUpdateProposal, proposal_id)
    if not prop: raise HTTPException(404)

    # 如果上传时选择了"替换模式"，先清空现有实习/项目/技能
    if prop.clear_existing:
        from backend.src.models.profile import Internship, Project, Skill
        from sqlmodel import delete as _del
        for model in [Internship, Project, Skill]:
            session.exec(_del(model).where(model.profile_id == prop.profile_id))
        session.commit()
        logger.info("替换模式: 已清空 profile_id=%d 的实习/项目/技能", prop.profile_id)

    changes = json.loads(prop.changes or "[]")
    applied = 0
    for ch in changes:
        if ch.get("id") in accepted_ids:
            _apply_change(ch, session, profile_id=prop.profile_id)
            applied += 1

    prop.status = "confirmed"
    session.add(prop)
    session.commit()
    return {"applied": applied}


async def _llm_parse_resume(text: str) -> list[dict]:
    """使用 LLM 从简历文本提取结构化更新建议。"""
    import json as _json
    from backend.src.llm.client import llm_client

    prompt = f"""你是一个简历解析器。从以下文本中提取所有可以结构化表示的信息。返回严格 JSON（只返回 JSON，不要解释）：

{{
  "name": "姓名（提取不到填 null）",
  "phone": "电话",
  "email": "邮箱",
  "internships": [
    {{"company": "公司全称", "position": "职位", "start_date": "起始如2024.06", "end_date": "结束或至今", "achievements": ["量化成果1", "成果2"]}}
  ],
  "projects": [
    {{"name": "项目名称", "role": "你的角色", "tech_stack": ["技术1","技术2"], "challenge": "核心挑战", "solution": "你的方案", "result": "量化结果"}}
  ],
  "skills": [
    {{"category": "language|framework|tool|other", "name": "技能名（规范化如K8s→Kubernetes）", "proficiency": "了解|熟悉|精通"}}
  ]
}}

提取不到的字段返回空数组或 null。只返回 JSON，不要任何其他文本。

简历文本：
{text[:8000]}"""

    try:
        resp = await llm_client.chat(
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
        )
        resp = resp.strip()
        if resp.startswith("```"):
            resp = resp.split("\n", 1)[-1].rsplit("```", 1)[0]
        data = _json.loads(resp)
    except _json.JSONDecodeError:
        logger.warning("LLM 简历解析返回了无效 JSON，原始响应前 200 字符: %s", resp[:200] if 'resp' in dir() else "N/A")
        return []
    # 注意：llm_client.chat() 的网络/API 异常不在此捕获，向上传播到 upload_document 处理

    proposals = []
    idx = 0

    # Profile fields
    for key in ("name", "phone", "email"):
        val = data.get(key)
        if val and str(val) not in ("null", "未知", "None", ""):
            idx += 1
            proposals.append({"id": f"c{idx}", "op": "update", "target": "profile",
                              "value": {key: str(val)}, "conflict": False})

    # Internships
    for item in data.get("internships", []) or []:
        if item.get("company") and item.get("company") not in ("null", "未知"):
            idx += 1
            proposals.append({"id": f"c{idx}", "op": "add", "target": "internship",
                              "value": item, "conflict": False})

    # Projects
    for item in data.get("projects", []) or []:
        if item.get("name") and item.get("name") not in ("null", "未知"):
            idx += 1
            proposals.append({"id": f"c{idx}", "op": "add", "target": "project",
                              "value": item, "conflict": False})

    # Skills
    for item in data.get("skills", []) or []:
        if item.get("name") and item.get("name") not in ("null", "未知"):
            idx += 1
            proposals.append({"id": f"c{idx}", "op": "add", "target": "skill",
                              "value": item, "conflict": False})

    return proposals


def _apply_change(ch: dict, session: Session, profile_id: int):
    target = ch.get("target", "")
    value = ch.get("value", {})
    if target == "profile":
        from backend.src.services.profile_service import update_profile
        update_profile(session, value)
    elif target == "internship":
        from backend.src.services.profile_service import create_internship
        create_internship(session, profile_id, value)
    elif target == "project":
        from backend.src.services.profile_service import create_project
        create_project(session, profile_id, value)
    elif target == "skill":
        from backend.src.services.profile_service import create_skill
        create_skill(session, profile_id, value)
